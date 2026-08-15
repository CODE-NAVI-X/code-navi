"""Application rules between the browser API and Piston adapter."""

from __future__ import annotations

import base64
import binascii
import io
import re
import sqlite3
from dataclasses import dataclass
from threading import Lock
from time import monotonic, perf_counter
from typing import Any, Protocol
from uuid import UUID, uuid4

from .ai_evaluation import (
    AiEvaluationError,
    AiEvaluator,
    AiTutor,
    PracticeSetPlanner,
    ProblemOrganizer,
)
from .config import Settings
from .evaluation import AiFeedback, RuleAssessment, classify_execution
from .judging import JudgeResult, judge_submission
from .learning_records import LearningRecordStore
from .piston import ExecutionLimits, ExecutionResult, PistonError, RuntimeInfo
from .problem_imports import ImportedProblem, analyze_problem_text
from .problem_sets import (
    MAX_PRACTICE_SET_SIZE,
    PracticeSetProblem,
    PracticeSetResult,
    apply_practice_set_plan,
    build_practice_set,
)
from .problems.catalog import build_default_problem_repository
from .problems.repository import ProblemRepository


class ValidationError(ValueError):
    """Raised when a browser execution request violates the public contract."""


class PistonGateway(Protocol):
    """Execution capabilities required by the web application."""

    def list_runtimes(self) -> tuple[RuntimeInfo, ...]:
        """Return installed runtimes."""

    def execute_python(
        self,
        source: str,
        stdin: str,
        *,
        version: str,
        limits: ExecutionLimits,
    ) -> ExecutionResult:
        """Execute Python source once."""


@dataclass(frozen=True)
class ApiResponse:
    """HTTP-independent API result used by the request handler and tests."""

    status_code: int
    body: dict[str, Any]


@dataclass(frozen=True, slots=True)
class PendingEvaluation:
    """Server-owned execution context referenced by an opaque browser ticket."""

    source: str
    result: ExecutionResult
    assessment: RuleAssessment
    learner_id: str | None
    record_id: str | None
    created_at: float


@dataclass(frozen=True, slots=True)
class PendingSubmission:
    source: str
    problem_id: str
    version: int
    description: str
    result: JudgeResult
    learner_id: str | None
    created_at: float


PENDING_EVALUATION_TTL_SECONDS = 300.0
MAX_PENDING_EVALUATIONS = 256
PENDING_SUBMISSION_TTL_SECONDS = 900.0
MAX_PENDING_SUBMISSIONS = 256
MAX_PROBLEM_IMPORT_TEXT_BYTES = 64 * 1024
MAX_PROBLEM_IMPORT_FILE_BYTES = 2 * 1024 * 1024
MAX_UPLOADED_PROBLEM_TEXT_BYTES = 8 * 1024
MAX_UPLOADED_PROBLEM_HINT_BYTES = 1 * 1024
MAX_UPLOADED_PROBLEM_ID_BYTES = 160
MAX_UPLOADED_PROBLEM_TAG_BYTES = 64
MAX_UPLOADED_PROBLEM_SAMPLE_BYTES = 4 * 1024


def _problem_organization_changed(
    original: list[ImportedProblem], organized: list[ImportedProblem]
) -> bool:
    return [
        (problem.import_id, problem.difficulty, problem.tags, problem.order_reason)
        for problem in original
    ] != [
        (problem.import_id, problem.difficulty, problem.tags, problem.order_reason)
        for problem in organized
    ]


def _filename_suffix(filename: str | None) -> str:
    return (filename or "").lower().rsplit(".", 1)[-1] if filename and "." in filename else ""


def _decode_problem_import_file(content_base64: str, filename: str | None) -> str:
    try:
        raw = base64.b64decode(content_base64, validate=True)
    except (binascii.Error, ValueError) as error:
        raise ValidationError("contentBase64 must be valid base64") from error
    if len(raw) > MAX_PROBLEM_IMPORT_FILE_BYTES:
        raise ValidationError("uploaded file exceeds the 2 MiB import limit")

    suffix = _filename_suffix(filename)
    if suffix == "docx":
        return _extract_docx_text(raw)
    if suffix == "pdf":
        return _extract_pdf_text(raw)
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError as error:
        raise ValidationError("uploaded file must be UTF-8 text, DOCX, or PDF") from error


def _extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document

        document = Document(io.BytesIO(raw))
    except Exception as error:
        raise ValidationError("could not read DOCX problem file") from error

    lines: list[str] = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" ".join(cells))
    return "\n".join(line for line in lines if line).strip()


def _extract_pdf_text(raw: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        lines = [_normalize_pdf_text(page.extract_text() or "") for page in reader.pages]
    except Exception as error:
        raise ValidationError("could not read PDF problem file") from error
    return "\n".join(line for line in lines if line).strip()


def _normalize_pdf_text(text: str) -> str:
    compact = re.sub(r"[ \t]+", " ", text).strip()
    return re.sub(
        r"(?<!^)(?=(?:题目|练习|Problem|Exercise|描述|Description|输入|Input|输出|Output)\s*[：:])",
        "\n",
        compact,
        flags=re.IGNORECASE,
    ).strip()


def _practice_set_candidate(problem: PracticeSetProblem) -> dict[str, object]:
    return {
        "id": problem.id,
        "source": problem.source,
        "title": problem.title,
        "description": problem.description,
        "difficulty": problem.difficulty,
        "tags": list(problem.tags),
        "judgeable": problem.judgeable,
        "generationReason": problem.generation_reason,
        "limitations": list(problem.limitations),
    }


class CompilerApplication:
    """Validate public requests and expose stable compiler responses."""

    def __init__(
        self,
        gateway: PistonGateway,
        settings: Settings,
        *,
        evaluator: AiEvaluator | None = None,
        tutor: AiTutor | None = None,
        organizer: ProblemOrganizer | None = None,
        practice_set_planner: PracticeSetPlanner | None = None,
        ai_status: str = "disabled",
        ai_message: str = "未配置 AI 模型，规则识别仍可使用。",
        record_store: LearningRecordStore | None = None,
        problem_repository: ProblemRepository | None = None,
    ) -> None:
        self._gateway = gateway
        self._settings = settings
        self._evaluator = evaluator
        self._tutor = tutor
        self._organizer = organizer
        self._practice_set_planner = practice_set_planner
        self._ai_status = ai_status
        self._ai_message = ai_message
        self._record_store = record_store
        self._problem_repository = problem_repository or build_default_problem_repository()
        self._pending_evaluations: dict[str, PendingEvaluation] = {}
        self._pending_lock = Lock()
        self._submissions: dict[str, PendingSubmission] = {}
        self._submission_lock = Lock()
        self._limits = ExecutionLimits(
            wall_time_ms=settings.run_timeout_ms,
            cpu_time_ms=settings.run_cpu_time_ms,
            memory_bytes=settings.run_memory_limit_bytes,
            output_bytes=settings.output_limit_bytes,
        )

    def runtime_status(self) -> ApiResponse:
        """Report whether the pinned Python runtime is installed."""

        try:
            runtimes = self._gateway.list_runtimes()
        except PistonError:
            return ApiResponse(
                503,
                {
                    "ready": False,
                    "message": "执行服务暂时不可用，请稍后重试。",
                },
            )

        ready = any(
            runtime.version == self._settings.python_version
            and "python" in {runtime.language, *runtime.aliases}
            for runtime in runtimes
        )
        return ApiResponse(
            200 if ready else 503,
            {
                "ready": ready,
                "language": "Python",
                "version": self._settings.python_version,
                "limits": {
                    "wallTimeMs": self._limits.wall_time_ms,
                    "memoryBytes": self._limits.memory_bytes,
                    "sourceBytes": self._settings.max_source_bytes,
                },
                "message": "运行环境已就绪" if ready else "Python 运行时尚未安装",
                "ai": {
                    "status": self._ai_status,
                    "message": self._ai_message,
                },
            },
        )

    def execute(self, payload: Any) -> ApiResponse:
        """Validate a browser request, execute Python and normalize failures."""

        try:
            source, stdin, learner_id, ai_enabled = self._validate_payload(payload)
        except ValidationError as error:
            return ApiResponse(400, {"error": str(error)})

        execution_started = perf_counter()
        try:
            result = self._gateway.execute_python(
                source,
                stdin,
                version=self._settings.python_version,
                limits=self._limits,
            )
        except PistonError:
            return ApiResponse(
                503,
                {"error": "执行服务暂时不可用，请确认 Piston 与 Python 运行时已经启动。"},
            )
        execution_round_trip_ms = round((perf_counter() - execution_started) * 1_000)

        assessment = classify_execution(result)
        should_defer_ai = (
            ai_enabled and self._evaluator is not None and assessment.category != "system_error"
        )
        if not ai_enabled:
            feedback = None
            ai_body: dict[str, Any] = {
                "status": "disabled",
                "message": "本次运行未开启 AI 评析。",
            }
        elif should_defer_ai:
            feedback = None
            ai_body = {
                "status": "pending",
                "message": "代码运行完成，AI 正在评析。",
            }
        else:
            feedback, ai_body = self._evaluate(source, result, assessment, learner_id)
        record_body: dict[str, Any] | None = None
        record_id: str | None = None
        if learner_id is not None and self._record_store is not None:
            try:
                record = self._record_store.add(
                    learner_id,
                    source,
                    result,
                    assessment,
                    ai_status=ai_body["status"],
                    feedback=feedback,
                )
                record_id = record.record_id
                record_body = record.as_dict()
            except (OSError, ValueError, sqlite3.Error):
                record_body = {"status": "unavailable"}

        if should_defer_ai:
            ai_body["evaluationId"] = self._queue_evaluation(
                source,
                result,
                assessment,
                learner_id,
                record_id,
            )

        body = result.as_dict()
        body.update(
            {
                "assessment": assessment.as_dict(),
                "ai": ai_body,
                "record": record_body,
                "serviceTiming": {"executionRoundTripMs": execution_round_trip_ms},
            }
        )
        return ApiResponse(200, body)

    def evaluate(self, payload: Any) -> ApiResponse:
        """Consume one execution ticket and produce optional AI feedback."""

        try:
            evaluation_id, learner_id = self._validate_evaluation_payload(payload)
        except ValidationError as error:
            return ApiResponse(400, {"error": str(error)})

        pending = self._take_evaluation(evaluation_id, learner_id)
        if pending is None:
            return ApiResponse(404, {"error": "评析任务不存在或已经过期。"})

        evaluation_started = perf_counter()
        feedback, ai_body = self._evaluate(
            pending.source,
            pending.result,
            pending.assessment,
            pending.learner_id,
        )
        evaluation_round_trip_ms = round((perf_counter() - evaluation_started) * 1_000)
        record_body: dict[str, Any] | None = None
        if (
            pending.record_id is not None
            and pending.learner_id is not None
            and self._record_store is not None
        ):
            try:
                record = self._record_store.update_feedback(
                    pending.record_id,
                    pending.learner_id,
                    ai_status=ai_body["status"],
                    feedback=feedback,
                )
                record_body = None if record is None else record.as_dict()
            except (OSError, ValueError, sqlite3.Error):
                record_body = {"status": "unavailable"}
        return ApiResponse(
            200,
            {
                "ai": ai_body,
                "record": record_body,
                "serviceTiming": {"evaluationRoundTripMs": evaluation_round_trip_ms},
            },
        )

    def submit(self, payload: Any) -> ApiResponse:
        """Judge a problem using tests and limits owned by the server."""
        try:
            problem_id, version, source, learner_id = self._validate_submit_payload(payload)
        except ValidationError as error:
            return ApiResponse(400, {"error": str(error)})
        problem = self._problem_repository.get(problem_id, version)
        if problem is None:
            return ApiResponse(404, {"error": "problem or version not found"})
        try:
            result = judge_submission(
                source, problem, self._gateway, self._settings.python_version, self._limits
            )
        except PistonError:
            return ApiResponse(503, {"error": "execution service is temporarily unavailable"})
        submission_id = str(uuid4())
        pending = PendingSubmission(
            source,
            problem.problem_id,
            problem.version,
            self._problem_description(problem.problem_id),
            result,
            learner_id,
            monotonic(),
        )
        with self._submission_lock:
            self._discard_expired_submissions(monotonic())
            while len(self._submissions) >= MAX_PENDING_SUBMISSIONS:
                self._submissions.pop(next(iter(self._submissions)))
            self._submissions[submission_id] = pending
        body = result.as_dict()
        body.update(
            {
                "submissionId": submission_id,
                "problemId": problem.problem_id,
                "problemVersion": problem.version,
            }
        )
        return ApiResponse(200, body)

    def guidance(self, payload: Any) -> ApiResponse:
        """Return bounded guidance from server-owned submission context."""
        try:
            submission_id, message, learner_id, history = self._validate_guidance_payload(payload)
        except ValidationError as error:
            return ApiResponse(400, {"error": str(error)})
        with self._submission_lock:
            self._discard_expired_submissions(monotonic())
            submission = self._submissions.get(submission_id)
        if submission is None or submission.learner_id != learner_id:
            return ApiResponse(404, {"error": "submission context not found or expired"})
        if self._tutor is None:
            return ApiResponse(503, {"error": "AI tutor is not available"})
        context = {
            "problemId": submission.problem_id,
            "problemVersion": submission.version,
            "problemDescription": submission.description,
            "source": submission.source[:20_000],
            "verdict": submission.result.verdict,
            "score": submission.result.score,
            "publicTests": [
                item.as_dict() for item in submission.result.test_results if not item.hidden
            ],
        }
        try:
            reply = self._tutor.chat(message, context, history, learner_id)
        except AiEvaluationError:
            return ApiResponse(503, {"error": "AI tutor is temporarily unavailable"})
        return ApiResponse(200, {"submissionId": submission_id, "ai": reply})

    def learning_records(self, learner_id: str | None) -> ApiResponse:
        """Return privacy-minimized records for one anonymous browser identity."""

        try:
            normalized_id = self._validate_learner_id(learner_id, required=True)
        except ValidationError as error:
            return ApiResponse(400, {"error": str(error)})
        if self._record_store is None:
            return ApiResponse(503, {"error": "学习记录服务暂时不可用。"})
        try:
            records = self._record_store.list_for(normalized_id)
        except (OSError, ValueError, sqlite3.Error):
            return ApiResponse(503, {"error": "学习记录服务暂时不可用。"})
        return ApiResponse(200, {"records": [record.as_dict() for record in records]})

    def analyze_problem_import(self, payload: Any) -> ApiResponse:
        """Convert uploaded exercise text into a reviewable ordered problem list."""

        try:
            text, filename, learner_id = self._validate_problem_import_payload(payload)
        except ValidationError as error:
            return ApiResponse(400, {"error": str(error)})
        problems = analyze_problem_text(text, filename=filename)
        organizer_warnings: list[str] = []
        source = "deterministic_rule"
        if self._organizer is not None and problems:
            try:
                organized, organizer_warnings = self._organizer.organize(problems, learner_id)
                if _problem_organization_changed(problems, organized):
                    source = "rules_with_ai_organization"
                problems = organized
            except AiEvaluationError:
                organizer_warnings = ["AI 整理暂不可用，当前结果来自规则解析。"]
        return ApiResponse(
            200,
            {
                "source": source,
                "problems": [problem.as_dict() for problem in problems],
                "warnings": (
                    organizer_warnings
                    if problems and organizer_warnings
                    else []
                    if problems
                    else ["未能从上传内容中识别到题目，请补充题目描述、输入和输出说明。"]
                ),
            },
        )

    def generate_problem_set(self, payload: Any) -> ApiResponse:
        """Create an ordered practice set from built-ins, uploads, and optional AI planning."""

        try:
            request = self._validate_problem_set_payload(payload)
        except ValidationError as error:
            return ApiResponse(400, {"error": str(error)})
        result = build_practice_set(
            prompt=request["prompt"],
            target_count=request["target_count"],
            difficulty_range=request["difficulty_range"],
            knowledge_tags=request["knowledge_tags"],
            include_uploaded=request["include_uploaded"],
            uploaded_problems=request["uploaded_problems"],
        )
        planner_warnings: list[str] = []
        if self._practice_set_planner is not None and result.ordered_problems:
            try:
                planner_payload = self._practice_set_planner.plan_practice_set(
                    {
                        "prompt": request["prompt"],
                        "targetCount": request["target_count"],
                        "difficultyRange": list(request["difficulty_range"]),
                        "knowledgeTags": list(request["knowledge_tags"]),
                        "includeUploadedProblems": request["include_uploaded"],
                    },
                    [_practice_set_candidate(problem) for problem in result.ordered_problems],
                    request["learner_id"],
                )
                result = apply_practice_set_plan(result, planner_payload)
            except (AiEvaluationError, ValueError):
                planner_warnings = ["AI 练习集规划暂不可用，当前结果来自规则生成。"]
        if planner_warnings:
            result = PracticeSetResult(
                source=result.source,
                ordered_problems=result.ordered_problems,
                rationale=result.rationale,
                coverage=result.coverage,
                warnings=tuple(dict.fromkeys([*result.warnings, *planner_warnings])),
            )
        return ApiResponse(200, result.as_dict())

    def _validate_problem_import_payload(self, payload: Any) -> tuple[str, str | None, str | None]:
        if not isinstance(payload, dict):
            raise ValidationError("request body must be a JSON object")
        filename = payload.get("filename")
        if filename is not None and (
            not isinstance(filename, str) or len(filename) > 160 or "\x00" in filename
        ):
            raise ValidationError("filename is invalid")
        text = payload.get("text")
        content_base64 = payload.get("contentBase64")
        if content_base64 is not None:
            if not isinstance(content_base64, str) or not content_base64.strip():
                raise ValidationError("contentBase64 must be a non-empty string")
            text = _decode_problem_import_file(content_base64, filename)
        if not isinstance(text, str) or not text.strip():
            raise ValidationError("text is required")
        if "\x00" in text:
            raise ValidationError("text must not contain null bytes")
        if len(text.encode("utf-8")) > MAX_PROBLEM_IMPORT_TEXT_BYTES:
            raise ValidationError("text exceeds the 64 KiB import limit")
        learner_id = payload.get("learnerId")
        if learner_id is not None and not isinstance(learner_id, str):
            raise ValidationError("learnerId must be a string")
        return text.strip(), filename, learner_id

    def _validate_problem_set_payload(self, payload: Any) -> dict[str, Any]:
        if not isinstance(payload, dict):
            raise ValidationError("request body must be a JSON object")
        prompt = payload.get("prompt")
        if not isinstance(prompt, str) or not prompt.strip():
            raise ValidationError("prompt is required")
        if "\x00" in prompt or len(prompt.encode("utf-8")) > 4_096:
            raise ValidationError("prompt is invalid")
        target_count = payload.get("targetCount", 5)
        if not isinstance(target_count, int) or isinstance(target_count, bool):
            raise ValidationError("targetCount must be an integer")
        target_count = min(MAX_PRACTICE_SET_SIZE, max(1, target_count))
        difficulty_range = self._validate_difficulty_range(payload.get("difficultyRange"))
        knowledge_tags = self._validate_string_list(payload.get("knowledgeTags"), "knowledgeTags")
        include_uploaded = payload.get("includeUploadedProblems", True)
        if not isinstance(include_uploaded, bool):
            raise ValidationError("includeUploadedProblems must be a boolean")
        uploaded_problems = payload.get("uploadedProblems", [])
        if not isinstance(uploaded_problems, list):
            raise ValidationError("uploadedProblems must be an array")
        validated_uploaded_problems = tuple(
            self._validate_uploaded_practice_problem(item, index)
            for index, item in enumerate(uploaded_problems[:MAX_PRACTICE_SET_SIZE])
        )
        learner_id = payload.get("learnerId")
        if learner_id is not None and not isinstance(learner_id, str):
            raise ValidationError("learnerId must be a string")
        return {
            "prompt": prompt.strip(),
            "target_count": target_count,
            "difficulty_range": difficulty_range,
            "knowledge_tags": tuple(knowledge_tags),
            "include_uploaded": include_uploaded,
            "uploaded_problems": validated_uploaded_problems,
            "learner_id": learner_id,
        }

    def _validate_uploaded_practice_problem(
        self, value: Any, index: int
    ) -> dict[str, object]:
        if not isinstance(value, dict):
            raise ValidationError(f"uploadedProblems[{index}] must be an object")
        problem: dict[str, object] = {}
        for field in ("id", "importId"):
            text = self._optional_bounded_text(
                value.get(field),
                f"uploadedProblems[{index}].{field}",
                MAX_UPLOADED_PROBLEM_ID_BYTES,
            )
            if text is not None:
                problem[field] = text
        for field in ("title", "description"):
            text = self._optional_bounded_text(
                value.get(field),
                f"uploadedProblems[{index}].{field}",
                MAX_UPLOADED_PROBLEM_TEXT_BYTES,
            )
            if text is not None:
                problem[field] = text
        difficulty = value.get("difficulty")
        if difficulty is not None:
            if difficulty not in {"easy", "medium", "hard"}:
                raise ValidationError(f"uploadedProblems[{index}].difficulty is invalid")
            problem["difficulty"] = difficulty
        tags = value.get("tags")
        if tags is not None:
            if not isinstance(tags, list):
                raise ValidationError(f"uploadedProblems[{index}].tags must be an array")
            problem["tags"] = [
                self._required_bounded_text(
                    tag,
                    f"uploadedProblems[{index}].tags[{tag_index}]",
                    MAX_UPLOADED_PROBLEM_TAG_BYTES,
                )
                for tag_index, tag in enumerate(tags[:8])
            ]
        source = self._optional_bounded_text(
            value.get("source"),
            f"uploadedProblems[{index}].source",
            self._settings.max_source_bytes,
        )
        if source is not None:
            problem["source"] = source
        for field in ("starterCode", "inputHint", "outputHint"):
            limit = (
                self._settings.max_source_bytes
                if field == "starterCode"
                else MAX_UPLOADED_PROBLEM_HINT_BYTES
            )
            text = self._optional_bounded_text(
                value.get(field),
                f"uploadedProblems[{index}].{field}",
                limit,
            )
            if text is not None:
                problem[field] = text
        sample_tests = value.get("sampleTests")
        if sample_tests is not None:
            if not isinstance(sample_tests, list):
                raise ValidationError(f"uploadedProblems[{index}].sampleTests must be an array")
            problem["sampleTests"] = [
                self._validate_uploaded_sample_test(sample, index, sample_index)
                for sample_index, sample in enumerate(sample_tests[:4])
            ]
        return problem

    def _validate_uploaded_sample_test(
        self, value: Any, problem_index: int, sample_index: int
    ) -> dict[str, str]:
        if not isinstance(value, dict):
            raise ValidationError(
                f"uploadedProblems[{problem_index}].sampleTests[{sample_index}] must be an object"
            )
        return {
            "stdin": self._required_bounded_text(
                value.get("stdin"),
                f"uploadedProblems[{problem_index}].sampleTests[{sample_index}].stdin",
                MAX_UPLOADED_PROBLEM_SAMPLE_BYTES,
            ),
            "expectedOutput": self._required_bounded_text(
                value.get("expectedOutput"),
                f"uploadedProblems[{problem_index}].sampleTests[{sample_index}].expectedOutput",
                MAX_UPLOADED_PROBLEM_SAMPLE_BYTES,
            ),
        }

    @staticmethod
    def _optional_bounded_text(value: Any, field: str, limit: int) -> str | None:
        if value is None:
            return None
        return CompilerApplication._required_bounded_text(value, field, limit)

    @staticmethod
    def _required_bounded_text(value: Any, field: str, limit: int) -> str:
        if not isinstance(value, str):
            raise ValidationError(f"{field} must be a string")
        text = value.strip()
        if "\x00" in text or len(text.encode("utf-8")) > limit:
            raise ValidationError(f"{field} exceeds the server limit")
        return text

    def _validate_difficulty_range(self, value: Any) -> tuple[str, str]:
        if value is None:
            return ("easy", "hard")
        if (
            not isinstance(value, list)
            or len(value) != 2
            or any(item not in {"easy", "medium", "hard"} for item in value)
        ):
            raise ValidationError("difficultyRange must contain two difficulty values")
        order = {"easy": 0, "medium": 1, "hard": 2}
        low, high = value
        return (low, high) if order[low] <= order[high] else (high, low)

    def _validate_string_list(self, value: Any, field: str) -> list[str]:
        if value is None:
            return []
        if not isinstance(value, list):
            raise ValidationError(f"{field} must be an array")
        items = [item.strip() for item in value if isinstance(item, str) and item.strip()]
        return items[:8]

    def _validate_submit_payload(self, payload: Any) -> tuple[str, int, str, str | None]:
        if not isinstance(payload, dict):
            raise ValidationError("request body must be a JSON object")
        problem_id = payload.get("problemId")
        version = payload.get("problemVersion", 1)
        source = payload.get("source")
        if not isinstance(problem_id, str) or not problem_id.strip():
            raise ValidationError("problemId is required")
        if not isinstance(version, int) or isinstance(version, bool) or version <= 0:
            raise ValidationError("problemVersion must be a positive integer")
        if not isinstance(source, str) or not source.strip():
            raise ValidationError("source must be non-empty Python code")
        if "\x00" in source or len(source.encode("utf-8")) > self._settings.max_source_bytes:
            raise ValidationError("source exceeds the server limit")
        return (
            problem_id.strip(),
            version,
            source,
            self._validate_learner_id(payload.get("learnerId"), required=False),
        )

    def _validate_guidance_payload(
        self, payload: Any
    ) -> tuple[str, str, str | None, list[dict[str, str]]]:
        if not isinstance(payload, dict):
            raise ValidationError("request body must be a JSON object")
        submission_id = payload.get("submissionId")
        message = payload.get("message")
        if (
            not isinstance(submission_id, str)
            or not isinstance(message, str)
            or not message.strip()
        ):
            raise ValidationError("submissionId and message are required")
        try:
            parsed = UUID(submission_id)
        except ValueError as exc:
            raise ValidationError("submissionId must be a UUID") from exc
        if parsed.version != 4 or len(message) > 800:
            raise ValidationError("invalid submissionId or message length")
        learner_id = self._validate_learner_id(payload.get("learnerId"), required=False)
        raw_history = payload.get("history", [])
        if not isinstance(raw_history, list) or len(raw_history) > 8:
            raise ValidationError("history is invalid")
        history: list[dict[str, str]] = []
        for item in raw_history:
            if (
                not isinstance(item, dict)
                or item.get("role") not in {"user", "assistant"}
                or not isinstance(item.get("content"), str)
            ):
                raise ValidationError("history is invalid")
            history.append({"role": item["role"], "content": item["content"][:800]})
        return str(parsed), message.strip(), learner_id, history

    def _problem_description(self, problem_id: str) -> str:
        from .problems.catalog import DEFAULT_PROBLEM_DEFINITIONS

        for definition in DEFAULT_PROBLEM_DEFINITIONS:
            if definition.problem_id == problem_id:
                return definition.description
        return "完成题目要求并通过测试。"

    def _discard_expired_submissions(self, now: float) -> None:
        expired = [
            key
            for key, item in self._submissions.items()
            if now - item.created_at > PENDING_SUBMISSION_TTL_SECONDS
        ]
        for key in expired:
            self._submissions.pop(key, None)

    def _validate_payload(self, payload: Any) -> tuple[str, str, str | None, bool]:
        if not isinstance(payload, dict):
            raise ValidationError("请求体必须是 JSON 对象。")

        language = payload.get("language", "python")
        if language != "python":
            raise ValidationError("当前版本只支持 Python。")

        source = payload.get("source")
        stdin = payload.get("stdin", "")
        ai_enabled = payload.get("enableAi", True)
        learner_id = self._validate_learner_id(payload.get("learnerId"), required=False)
        if not isinstance(source, str) or not source.strip():
            raise ValidationError("请输入需要运行的 Python 代码。")
        if not isinstance(stdin, str):
            raise ValidationError("标准输入必须是文本。")
        if not isinstance(ai_enabled, bool):
            raise ValidationError("enableAi 必须是布尔值。")
        if "\x00" in source or "\x00" in stdin:
            raise ValidationError("代码和标准输入不能包含空字节。")
        if len(source.encode("utf-8")) > self._settings.max_source_bytes:
            raise ValidationError(f"代码不能超过 {self._settings.max_source_bytes // 1024} KiB。")
        if len(stdin.encode("utf-8")) > self._settings.max_stdin_bytes:
            raise ValidationError(
                f"标准输入不能超过 {self._settings.max_stdin_bytes // 1024} KiB。"
            )
        return source, stdin, learner_id, ai_enabled

    def _validate_evaluation_payload(self, payload: Any) -> tuple[str, str | None]:
        if not isinstance(payload, dict):
            raise ValidationError("请求体必须是 JSON 对象。")
        evaluation_id = payload.get("evaluationId")
        if not isinstance(evaluation_id, str):
            raise ValidationError("evaluationId 必须是 UUID 字符串。")
        try:
            parsed = UUID(evaluation_id)
        except ValueError as exc:
            raise ValidationError("evaluationId 必须是有效的 UUID。") from exc
        if parsed.version != 4:
            raise ValidationError("evaluationId 必须是 UUID v4。")
        learner_id = self._validate_learner_id(payload.get("learnerId"), required=False)
        return str(parsed), learner_id

    def _queue_evaluation(
        self,
        source: str,
        result: ExecutionResult,
        assessment: RuleAssessment,
        learner_id: str | None,
        record_id: str | None,
    ) -> str:
        evaluation_id = str(uuid4())
        now = monotonic()
        pending = PendingEvaluation(
            source,
            result,
            assessment,
            learner_id,
            record_id,
            now,
        )
        with self._pending_lock:
            self._discard_expired_evaluations(now)
            while len(self._pending_evaluations) >= MAX_PENDING_EVALUATIONS:
                oldest_id = next(iter(self._pending_evaluations))
                self._pending_evaluations.pop(oldest_id)
            self._pending_evaluations[evaluation_id] = pending
        return evaluation_id

    def _take_evaluation(
        self, evaluation_id: str, learner_id: str | None
    ) -> PendingEvaluation | None:
        with self._pending_lock:
            self._discard_expired_evaluations(monotonic())
            pending = self._pending_evaluations.get(evaluation_id)
            if pending is None or pending.learner_id != learner_id:
                return None
            return self._pending_evaluations.pop(evaluation_id)

    def _discard_expired_evaluations(self, now: float) -> None:
        expired = [
            evaluation_id
            for evaluation_id, pending in self._pending_evaluations.items()
            if now - pending.created_at > PENDING_EVALUATION_TTL_SECONDS
        ]
        for evaluation_id in expired:
            self._pending_evaluations.pop(evaluation_id, None)

    def _evaluate(
        self,
        source: str,
        result: ExecutionResult,
        assessment: RuleAssessment,
        learner_id: str | None,
    ) -> tuple[AiFeedback | None, dict[str, Any]]:
        if assessment.category == "system_error":
            return None, {
                "status": "skipped",
                "message": "执行环境异常，不对学生代码生成 AI 评价。",
            }
        if self._evaluator is None:
            return None, {"status": self._ai_status, "message": self._ai_message}
        try:
            feedback = self._evaluator.evaluate(source, result, assessment, learner_id)
        except AiEvaluationError:
            return None, {
                "status": "unavailable",
                "message": "AI 反馈暂时不可用；规则结论和学习记录不受影响。",
            }
        return feedback, feedback.as_dict()

    @staticmethod
    def _validate_learner_id(value: Any, *, required: bool) -> str | None:
        if value is None and not required:
            return None
        if not isinstance(value, str):
            raise ValidationError("learnerId 必须是 UUID 字符串。")
        try:
            parsed = UUID(value)
        except ValueError as exc:
            raise ValidationError("learnerId 必须是有效的 UUID。") from exc
        if parsed.version != 4:
            raise ValidationError("learnerId 必须是随机生成的 UUID v4。")
        return str(parsed)
