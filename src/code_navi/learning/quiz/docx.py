"""Word (.docx) exam-paper exporter for generated quizzes.

The layout follows the local "数学试卷排版大师" skill (``SKILL.md``) and the
observed structure of its sample paper (``集合高一数学阶段性测试卷1.docx``):

- A4, 1.27 cm (720 twips) margins on all sides.
- Title: Microsoft YaHei 14 pt bold, centered.
- Section header (一、单项选择题 / 二、填空题 / 三、解答题): SimSun 14 pt bold.
- Body: SimSun 12 pt with Times New Roman for Latin/numbers, 1.25 line spacing.
- Question stems are flush-left (顶格); options and short-answer sub-parts are
  indented 2 characters (``w:leftChars=200``).
- Four options on one line aligned via tab stops at 4.6 / 9.2 / 13.8 cm.
- LaTeX math (``$...$``) is converted to native Word equations (OMML) through
  ``latex2mathml`` + the bundled ``MML2OMML.XSL`` — matching the sample paper's
  OMML equations rather than static images or plain text.
- Question numbers are continuous across the whole paper.

The exporter is pure ``python-docx`` (no Pandoc, no Word COM), so it runs on
the Linux/Docker backend as well as on Windows.
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path

import latex2mathml.converter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

from .schemas import QuizQuestion

logger = logging.getLogger(__name__)

# ruff: noqa: E501 -- a few long format strings below read better unwrapped.

# ---------------------------------------------------------------------------
# Constants (aligned with the typesetting skill)
# ---------------------------------------------------------------------------

MARGIN_CM = 1.27
TITLE_FONT = "微软雅黑"
TITLE_SIZE = Pt(14)
SECTION_FONT = "宋体"
SECTION_SIZE = Pt(14)
BODY_FONT = "宋体"
BODY_LATIN = "Times New Roman"
BODY_SIZE = Pt(12)
LINE_SPACING = 1.25

# Tab stops for a 1x4 option row (cm): page width 21 - 2*1.27 = 18.46 usable,
# the skill's A4 / narrow-margin choice keeps B/C/D at 4.6/9.2/13.8 cm.
OPTION_TAB_STOPS_CM = (4.6, 9.2, 13.8)
OPTION_TAB_2COL_CM = 9.2

# Long-option fallback thresholds (best-effort width estimation).
_OPTION_LONG_CHARS = 14  # 4 options no longer fit one row → 2x2
_OPTION_VERY_LONG_CHARS = 26  # 2 options no longer fit a row → 4x1

_LATEX_INLINE_RE = re.compile(r"\$([^$\n]+)\$")

_SECTION_HEADERS = {
    "single": "一、单项选择题",
    "fill_blank": "二、填空题",
    "short_answer": "三、解答题",
}

#: Bundled copy of Word's MathML→OMML stylesheet, so formula rendering works
#: on servers without Microsoft Office installed.
_XSLT_PATH = Path(__file__).parent / "MML2OMML.XSL"


# ---------------------------------------------------------------------------
# OMML helpers
# ---------------------------------------------------------------------------


def _omml_from_latex(latex: str) -> etree._Element | None:
    """Convert one LaTeX snippet to an OMML ``m:oMath`` element.

    Returns ``None`` (caller falls back to literal text) on any conversion
    error so one bad formula never breaks the whole paper.
    """
    try:
        mml = latex2mathml.converter.convert(latex)
        xslt = etree.XSLT(etree.parse(str(_XSLT_PATH)))
        mmldoc = etree.fromstring(mml.encode("utf-8"))
        return etree.fromstring(etree.tostring(xslt(mmldoc)))
    except Exception as exc:  # noqa: BLE001 — degrade gracefully per formula
        logger.warning("OMML conversion failed for %r (%s); using literal text.", latex, exc)
        return None


# ---------------------------------------------------------------------------
# Low-level paragraph / run builders
# ---------------------------------------------------------------------------


def _set_font(
    run,
    size_pt: float = 12,
    bold: bool = False,
    east_asia: str = BODY_FONT,
    latin: str = BODY_LATIN,
) -> None:
    """Apply the skill's typography to one run."""
    run.font.name = latin
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asia)


def _char_indent(p, left_chars: int = 0, first_chars: int = 0) -> None:
    """Set char-based (w:ind) indentation used by the skill.

    ``left_chars`` indents the whole paragraph (options, sub-parts);
    ``first_chars`` indents only the first line (not used by the exam layout).
    """
    ppr = p._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if left_chars:
        ind.set(qn("w:leftChars"), str(left_chars * 100))
        ind.set(qn("w:left"), str(left_chars * 240))  # 240 twips/char @12pt fallback
    if first_chars:
        ind.set(qn("w:firstLineChars"), str(first_chars * 100))
        ind.set(qn("w:firstLine"), str(first_chars * 240))


def _add_tab_stops(p, positions_cm: tuple[float, ...]) -> None:
    for pos in positions_cm:
        p.paragraph_format.tab_stops.add_tab_stop(Cm(pos), WD_TAB_ALIGNMENT.LEFT)


def _add_rich_text(p, text: str, *, size_pt: float = 12, bold: bool = False) -> None:
    """Add runs to ``p``, converting each ``$...$`` segment to an OMML equation."""
    pos = 0
    for match in _LATEX_INLINE_RE.finditer(text):
        if match.start() > pos:
            _add_plain_run(p, text[pos : match.start()], size_pt=size_pt, bold=bold)
        omml = _omml_from_latex(match.group(1))
        if omml is not None:
            p._p.append(omml)
        else:
            _add_plain_run(p, f"${match.group(1)}$", size_pt=size_pt, bold=bold)
        pos = match.end()
    if pos < len(text):
        _add_plain_run(p, text[pos:], size_pt=size_pt, bold=bold)


def _add_plain_run(p, text: str, *, size_pt: float = 12, bold: bool = False) -> None:
    run = p.add_run(text)
    _set_font(run, size_pt=size_pt, bold=bold)


def _paragraph(doc: Document, *, align=None) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    if align is not None:
        pf.alignment = align
    return p


# ---------------------------------------------------------------------------
# Page / style setup
# ---------------------------------------------------------------------------


def _configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)

    # Make every paragraph inherit the body typography.
    style = doc.styles["Normal"]
    style.font.name = BODY_LATIN
    style.font.size = BODY_SIZE
    style.font.color.rgb = RGBColor(0, 0, 0)
    rfonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), BODY_FONT)
    style.paragraph_format.line_spacing = LINE_SPACING


def _add_title(doc: Document, text: str) -> None:
    p = _paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    _set_font(run, size_pt=14, bold=True, east_asia=TITLE_FONT, latin=TITLE_FONT)


def _add_meta(doc: Document, text: str) -> None:
    p = _paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_rich_text(p, text)


def _add_section_header(doc: Document, text: str) -> None:
    p = _paragraph(doc)
    run = p.add_run(text)
    _set_font(run, size_pt=14, bold=True, east_asia=SECTION_FONT, latin=SECTION_FONT)


def _add_blank_lines(doc: Document, count: int) -> None:
    # A literal ``&nbsp;`` would be written as the six visible characters
    # ``&nbsp;`` into the run (python-docx does not decode HTML entities).
    # Use the real Unicode no-break space so Word keeps the line height but
    # renders an empty writing line.
    for _ in range(count):
        p = _paragraph(doc)
        _add_plain_run(p, " ")  # U+00A0 NO-BREAK SPACE — keeps the line


# ---------------------------------------------------------------------------
# Question renderers
# ---------------------------------------------------------------------------


def _render_choice_options(doc: Document, q: QuizQuestion) -> None:
    options = [o for o in (q.options or []) if o.label.strip()]
    if not options:
        return
    max_len = max(len(o.label) for o in options)
    p = _paragraph(doc)
    _char_indent(p, left_chars=2)

    if max_len >= _OPTION_VERY_LONG_CHARS:
        # 4x1: each option on its own indented line.
        for option in options:
            line = _paragraph(doc)
            _char_indent(line, left_chars=2)
            _add_plain_run(line, f"{option.value}. ", size_pt=12)
            _add_rich_text(line, option.label)
    elif max_len >= _OPTION_LONG_CHARS:
        # 2x2: A/B on the first line, C/D on the second, aligned at 9.2 cm.
        _add_tab_stops(p, (OPTION_TAB_2COL_CM,))
        _render_option_row(p, options[:2])
        p2 = _paragraph(doc)
        _char_indent(p2, left_chars=2)
        _add_tab_stops(p2, (OPTION_TAB_2COL_CM,))
        _render_option_row(p2, options[2:])
    else:
        # 1x4: A/B/C/D on one line aligned at 4.6/9.2/13.8 cm.
        _add_tab_stops(p, OPTION_TAB_STOPS_CM)
        _render_option_row(p, options)


def _render_option_row(p, options) -> None:
    for i, option in enumerate(options):
        if i:
            p.add_run().add_tab()
        _add_plain_run(p, f"{option.value}. ", size_pt=12)
        _add_rich_text(p, option.label)


def _render_question(doc: Document, number: int, q: QuizQuestion) -> None:
    p = _paragraph(doc)
    _add_plain_run(p, f"{number}. ", size_pt=12)
    _add_rich_text(p, q.question)

    if q.type == "single":
        _render_choice_options(doc, q)
    elif q.type == "fill_blank":
        return  # blank underscores are already part of the stem
    else:  # short_answer
        # 7-8 blank lines give students room to write, per the skill.
        _add_blank_lines(doc, 7)


def _render_answer_section(doc: Document, questions: list[QuizQuestion]) -> None:
    doc.add_page_break()
    _add_title(doc, "参考答案")

    last_section = None
    number = 0
    for q in questions:
        if q.type != last_section:
            last_section = q.type
            _add_section_header(doc, _SECTION_HEADERS[q.type])
        number += 1
        p = _paragraph(doc)
        _add_plain_run(p, f"{number}. ", size_pt=12)
        answer_text = _format_answer(q)
        if answer_text:
            _add_rich_text(p, answer_text)
        if q.analysis:
            p2 = _paragraph(doc)
            _char_indent(p2, left_chars=2)
            _add_plain_run(p2, "解析：", size_pt=12)
            _add_rich_text(p2, q.analysis)


def _format_answer(q: QuizQuestion) -> str:
    if q.type == "single" and q.answer:
        return "答案：" + "、".join(q.answer)
    if q.type == "fill_blank" and q.answer:
        return "答案：" + "，".join(q.answer)
    if q.type == "short_answer":
        return "参考答案：" + (q.analysis or "")
    return ""


# ---------------------------------------------------------------------------
# Public exporter
# ---------------------------------------------------------------------------


def _group_by_type(questions: list[QuizQuestion]) -> list[tuple[str, list[QuizQuestion]]]:
    """Preserve the fixed section order: single → fill_blank → short_answer."""
    order = ("single", "fill_blank", "short_answer")
    return [
        (t, [q for q in questions if q.type == t])
        for t in order
        if any(q.type == t for q in questions)
    ]


def export_quiz_docx(
    *,
    knowledge_point: str,
    questions: list[QuizQuestion],
    with_answer: bool = False,
) -> bytes:
    """Render the quiz as a Word exam paper and return the .docx bytes."""
    doc = Document()
    _configure_document(doc)

    _add_title(doc, f"《{knowledge_point}》练习题")
    total = sum(q.points for q in questions)
    _add_meta(doc, f"（共 {len(questions)} 题，满分 {total} 分）")

    number = 0
    for section, section_questions in _group_by_type(questions):
        _add_section_header(doc, _SECTION_HEADERS[section])
        for q in section_questions:
            number += 1
            _render_question(doc, number, q)

    if with_answer:
        _render_answer_section(doc, questions)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
